"""
Resume Parser Service
Extracts raw text from PDF and DOCX files.
"""
import io
from fastapi import UploadFile


async def extract_text(file: UploadFile) -> str:
    """Extract plain text from uploaded PDF or DOCX resume."""
    content = await file.read()
    filename = file.filename or ""

    if filename.lower().endswith(".pdf"):
        return _extract_pdf(content)
    elif filename.lower().endswith(".docx"):
        return _extract_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Only PDF and DOCX are supported.")


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text("text"))
        return "\n".join(text_parts).strip()
    except ImportError:
        raise RuntimeError("PyMuPDF is not installed. Run: pip install PyMuPDF")
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF: {str(e)}")


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {str(e)}")
